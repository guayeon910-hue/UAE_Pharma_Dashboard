"""싱가포르 보고서 템플릿 구조 분석 스크립트"""
import sys
import os
sys.stdout.reconfigure(encoding='utf-8')

import docx

files = [
    r'C:\Users\user\Documents\카카오톡 받은 파일\SG_최종보고서(2-3-1 합친버전).docx',
    r'C:\Users\user\Documents\카카오톡 받은 파일\SG_00_표지.docx',
    r'C:\Users\user\Documents\카카오톡 받은 파일\SG_01_시장보고서_Sereterol.docx',
    r'C:\Users\user\Documents\카카오톡 받은 파일\SG_02_수출가격전략_Gadvoa (1).docx',
    r'C:\Users\user\Documents\카카오톡 받은 파일\SG_03_바이어리스트.docx',
]

for fpath in files:
    fname = os.path.basename(fpath)
    if not os.path.exists(fpath):
        print(f'[MISSING] {fname}')
        continue
    doc = docx.Document(fpath)
    print(f'\n{"="*60}')
    print(f'FILE: {fname}')
    print(f'Sections: {len(doc.sections)}, Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}')
    print(f'{"="*60}')
    
    # Print all paragraphs with style info
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            style = p.style.name if p.style else 'None'
            text = p.text[:150]
            print(f'  P{i:03d} [{style}]: {text}')
    
    # Print table structure
    for ti, table in enumerate(doc.tables):
        print(f'\n  TABLE {ti}: {len(table.rows)} rows x {len(table.columns)} cols')
        for ri, row in enumerate(table.rows[:5]):
            cells = [cell.text[:40] for cell in row.cells]
            print(f'    Row {ri}: {cells}')
        if len(table.rows) > 5:
            print(f'    ... ({len(table.rows) - 5} more rows)')
