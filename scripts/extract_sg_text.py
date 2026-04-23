"""SG 보고서 템플릿 전체 텍스트 추출 (UTF-8 인코딩)"""
import sys, os, json
sys.stdout.reconfigure(encoding='utf-8')
import docx

files = [
    (r'C:\Users\user\Documents\카카오톡 받은 파일\SG_00_표지.docx', 'cover'),
    (r'C:\Users\user\Documents\카카오톡 받은 파일\SG_01_시장보고서_Sereterol.docx', 'market_report'),
    (r'C:\Users\user\Documents\카카오톡 받은 파일\SG_02_수출가격전략_Gadvoa (1).docx', 'price_strategy'),
    (r'C:\Users\user\Documents\카카오톡 받은 파일\SG_03_바이어리스트.docx', 'buyer_list'),
]

output = {}
for fpath, key in files:
    if not os.path.exists(fpath):
        output[key] = f'FILE NOT FOUND: {fpath}'
        continue
    doc = docx.Document(fpath)
    paragraphs = []
    for i, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style else 'None'
        paragraphs.append({'idx': i, 'style': style, 'text': p.text})
    
    tables = []
    for ti, table in enumerate(doc.tables):
        rows_data = []
        for ri, row in enumerate(table.rows):
            cells = [cell.text for cell in row.cells]
            rows_data.append(cells)
        tables.append({'table_idx': ti, 'rows': rows_data})
    
    output[key] = {
        'file': os.path.basename(fpath),
        'sections': len(doc.sections),
        'paragraph_count': len(doc.paragraphs),
        'table_count': len(doc.tables),
        'paragraphs': paragraphs,
        'tables': tables,
    }

with open(r'C:\Users\user\Desktop\UAE\scripts\sg_template_data.json', 'w', encoding='utf-8') as f:
    json.dump(output, f, ensure_ascii=False, indent=2)

print('Done! Saved to sg_template_data.json')
