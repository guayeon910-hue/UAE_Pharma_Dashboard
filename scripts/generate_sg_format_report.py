"""SG 양식 기반 UAE 보고서 DOCX 생성기.

SG 템플릿 4개(표지/시장보고서/수출가격전략/바이어리스트)를
UAE 데이터로 채워 DOCX 파일을 생성하고 PDF로 변환합니다.
"""
import sys, os, re, json, copy, asyncio, logging
from pathlib import Path
from datetime import datetime, timezone
from typing import Any

sys.stdout.reconfigure(encoding='utf-8')
ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

log = logging.getLogger("sg_report")

# ── UAE 거시환경 상수 ─────────────────────────────────────────────────────────
UAE_MACRO = {
    "country": "UAE(아랍에미리트)",
    "country_en": "United Arab Emirates",
    "population": "약 10,000,000명 (World Bank, 2024)",
    "gdp_per_capita": "USD 43,103 (IMF 2024)",
    "pharma_market": "USD 5.2 Billion (2024 추산, IQVIA)",
    "health_spend": "GDP 대비 약 4.5% (WHO, 2023)",
    "currency": "AED (디르함)",
    "regulator": "EDE (에미리트의약품청), MOHAP, DoH(아부다비), DHA(두바이)",
    "tax": "VAT 5%",
    "procurement": "Rafed/ADGPG 공공조달 플랫폼",
}


def _add_heading(doc, text, level=1):
    """스타일 없이 수동 포매팅으로 제목 추가."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 0:
        run.font.size = Pt(22)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = RGBColor(27, 42, 74)
    elif level == 2:
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = RGBColor(27, 42, 74)
    return p


def _add_body(doc, text, bold=False, size=10):
    """본문 텍스트 추가."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p


def _add_field(doc, label, value):
    """라벨 + 값 형태의 필드 라인."""
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}  ")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(str(value or "—"))
    r2.font.size = Pt(10)
    return p


# ── 00: 표지 ─────────────────────────────────────────────────────────────────

def generate_cover(output_path: Path, product_name: str = ""):
    doc = Document()
    for _ in range(12):
        doc.add_paragraph("")
    _add_heading(doc, "UAE(아랍에미리트) 진출 전략 보고서", level=0)
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("한국유나이티드제약")
    run.font.size = Pt(14)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(datetime.now().strftime("%Y-%m-%d"))
    run2.font.size = Pt(12)
    for _ in range(8):
        doc.add_paragraph("")
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("수출가격 전략 - 바이어 후보 리스트 - 시장분석")
    run3.font.size = Pt(11)
    doc.save(str(output_path))
    return output_path


# ── 01: 시장보고서 ────────────────────────────────────────────────────────────

def generate_market_report(output_path: Path, analysis: dict):
    doc = Document()
    trade = analysis.get("trade_name", "")
    inn = analysis.get("inn", "")
    
    _add_heading(doc, f"UAE 시장보고서 — {trade}", level=0)
    _add_body(doc, f"{trade} ({inn})  |  UAE  |  {datetime.now().strftime('%Y-%m-%d')}", size=10)
    
    # 1. 의료 거시환경
    _add_heading(doc, "1. 의료 거시환경 파악", level=1)
    _add_field(doc, "인구", UAE_MACRO["population"])
    _add_field(doc, "1인당 GDP", UAE_MACRO["gdp_per_capita"])
    _add_field(doc, "의약품 시장 규모", UAE_MACRO["pharma_market"])
    _add_field(doc, "보건 지출", UAE_MACRO["health_spend"])
    
    market_med = analysis.get("basis_market_medical", "")
    if market_med:
        doc.add_paragraph("")
        _add_body(doc, market_med)
    
    # 2. 무역/규제 환경
    _add_heading(doc, "2. 무역/규제 환경", level=1)
    
    _add_heading(doc, "▸ EDE 등록 현황", level=2)
    reg = analysis.get("basis_regulatory", "")
    _add_body(doc, reg or "EDE 등록 현황 확인 필요")
    
    _add_heading(doc, "▸ 진입 채널 권고", level=2)
    entry = analysis.get("entry_pathway", "")
    _add_body(doc, entry or "진입 경로 분석 필요")
    
    _add_heading(doc, "▸ 관세 및 무역", level=2)
    trade_basis = analysis.get("basis_trade", "")
    _add_body(doc, trade_basis or f"UAE는 의약품에 대해 관세 0%를 적용합니다(VAT {UAE_MACRO['tax']} 별도).")
    
    # 3. 참고 가격
    _add_heading(doc, "3. 참고 가격", level=1)
    price_pos = analysis.get("price_positioning", "")
    doh_ctx = analysis.get("doh_price_context", "")
    if doh_ctx:
        for line in doh_ctx.split("\n"):
            if line.strip():
                _add_body(doc, line.strip(), size=9)
    elif price_pos:
        _add_body(doc, price_pos)
    else:
        _add_body(doc, "DoH/DHA 참조 가격 데이터 미확보")
    
    # 4. 리스크 / 조건
    _add_heading(doc, "4. 리스크 / 조건", level=1)
    risks = analysis.get("risks_conditions", "")
    tatmeen = analysis.get("tatmeen_note", "")
    if risks:
        _add_body(doc, risks)
    if tatmeen:
        _add_heading(doc, "▸ Tatmeen 준수사항", level=2)
        _add_body(doc, tatmeen)
    
    # 5. 근거 및 출처
    _add_heading(doc, "5. 근거 및 출처", level=1)
    sources = analysis.get("sources", [])
    if sources:
        _add_heading(doc, "▸ 5-1. 참조 출처", level=2)
        for i, s in enumerate(sources, 1):
            name = s.get("name", "")
            url = s.get("url", "")
            _add_body(doc, f"No.{i}  {name}", bold=True, size=9)
            if url and url != "내부 데이터":
                _add_body(doc, f"출처: {url}", size=8)
    
    _add_heading(doc, "▸ 5-2. 사용된 DB/기관", level=2)
    db_list = [
        f"•  EDE Therapeutic Products Register — UAE 공식 의약품 등록 DB",
        f"•  DoH Abu Dhabi 참조 가격 리스트 — AED 약국공급가",
        f"•  DHA Dubai 약가표 — AED 기준 POM/OTC 분류",
        f"•  Aster Pharmacy 소매가 — UAE 민간 약국 실시간 소매가격",
        f"•  Rafed/ADGPG — UAE 공공 의료조달 플랫폼",
        f"•  Tatmeen GS1 — 의약품 추적 포털",
        f"•  Perplexity 실시간 규제·시장 정보",
    ]
    for item in db_list:
        _add_body(doc, item, size=9)
    
    doc.save(str(output_path))
    return output_path


# ── 02: 수출가격전략 ─────────────────────────────────────────────────────────

def generate_price_strategy(output_path: Path, analysis: dict, p2_data: dict = None):
    doc = Document()
    trade = analysis.get("trade_name", "")
    inn = analysis.get("inn", "")
    
    _add_heading(doc, f"UAE 수출 가격 전략 보고서 — {trade}", level=0)
    _add_body(doc, f"{trade} ({inn})  |  {datetime.now().strftime('%Y-%m-%d')}")
    
    # 1. UAE 거시 시장
    _add_heading(doc, "1. UAE 거시 시장", level=1)
    _add_body(doc, (
        f"UAE는 인구 약 1,000만 명의 고소득 산유국으로, 1인당 GDP USD 43,103(IMF 2024)로 "
        f"중동 최상위권입니다. 의료비 지출은 GDP의 약 4.5% 수준이며 의약품 시장 규모는 "
        f"약 USD 5.2B(2024 추산)으로 연평균 7~8% 성장세를 유지하고 있습니다. "
        f"DoH(아부다비)/DHA(두바이) 이원 규제 체계 하에서 EDE 등록이 시장 진입의 핵심 요건이며, "
        f"Rafed/ADGPG 공공조달 및 민간 약국 체인(Aster, Life 등)이 주요 유통 채널입니다."
    ))
    
    # 2. 단가 (시장 기준가)
    _add_heading(doc, f"2. {trade} 단가 (시장 기준가)", level=1)
    
    price_comp = analysis.get("price_comparison", {})
    competitors = price_comp.get("competitors", [])
    
    if competitors:
        # 0.01 이하 이상치 제외
        valid_prices = [c["price_aed"] for c in competitors if c.get("price_aed", 0) > 1]
        if valid_prices:
            avg_price = sum(valid_prices) / len(valid_prices)
            _add_field(doc, "기준 가격", f"AED {avg_price:.2f}")
        else:
            _add_field(doc, "기준 가격", "가격 데이터 수집 중")
    
    _add_field(doc, "산정 방식", "AI 분석 (Claude Haiku) + Aster Pharmacy 실시간 소매가 + DoH 참조가격")
    _add_field(doc, "시장 구분", "공공 / 민간")
    
    # 3. 거래처 참고 가격
    _add_heading(doc, "3. 거래처 참고 가격", level=1)
    if competitors:
        for c in competitors[:4]:
            source = c.get("source", "")
            tname = c.get("trade_name", "")
            price = c.get("price_aed", 0)
            if price > 1:
                _add_field(doc, "제품명", tname)
                _add_field(doc, "시장가", f"AED {price:.2f}  ({source})")
                doc.add_paragraph("")
    
    # 4. 가격 시나리오
    _add_heading(doc, "4. 가격 시나리오", level=1)
    
    if p2_data and p2_data.get("scenarios"):
        scenarios = p2_data["scenarios"]
        seg = p2_data.get("seg_label", "공공시장")
        formula = p2_data.get("formula_str", "")
        
        _add_heading(doc, f"▸ 4-1. {seg}", level=2)
        for sc in scenarios:
            label = sc.get("label", "")
            price = sc.get("price", 0)
            reason = sc.get("reason", "")
            doc.add_paragraph("")
            _add_field(doc, f"[{label}]", f"AED {price:.2f}" if price else "—")
            _add_body(doc, f"근거  {reason}", size=9)
            if formula:
                _add_body(doc, f"FOB 수출가 역산식  {formula}", size=9)
    elif competitors:
        valid = [c for c in competitors if c.get("price_aed", 0) > 1]
        if valid:
            prices = sorted([c["price_aed"] for c in valid])
            low = prices[0] * 0.75
            mid = sum(prices) / len(prices) * 0.85
            high = prices[-1] * 0.95
            
            for seg_name in ["공공 시장", "민간 시장"]:
                margin = 0.20 if "공공" in seg_name else 0.30
                _add_heading(doc, f"▸ {seg_name}", level=2)
                
                scenarios_gen = [
                    ("저가 진입", low, f"경쟁사 최저가 대비 75% 수준. 초기 시장 점유율 확보 전략."),
                    ("기준가", mid, f"경쟁사 평균가 대비 85% 수준. 품질 대비 합리적 가격 포지셔닝."),
                    ("프리미엄", high, f"경쟁사 최고가 대비 95% 수준. 품질·안전성 차별화 전략."),
                ]
                for label, price, reason in scenarios_gen:
                    doc.add_paragraph("")
                    _add_field(doc, f"[{label}]", f"AED {price:.2f}")
                    _add_body(doc, f"근거  {reason}", size=9)
                    fob = price / (1 + 0.05) / (1 + margin) / (1 + 0.15) * 0.27
                    _add_body(doc, f"FOB 수출가 역산식  AED {price:.0f} ÷ (1+VAT 5%) ÷ (1+마진 {int(margin*100)}%) ÷ (1+수입·유통비 15%) × 환율(0.27) ≈ USD {fob:.1f}", size=9)
    
    doc.add_paragraph("")
    _add_body(doc, "※ 본 산출 결과는 AI 분석에 기반한 추정치이므로, 최종 의사결정 전 반드시 담당자의 검토 및 확인이 필요합니다.", size=9)
    
    doc.save(str(output_path))
    return output_path


# ── 03: 바이어리스트 ─────────────────────────────────────────────────────────

def generate_buyer_list(output_path: Path, analysis: dict, buyers: list = None):
    doc = Document()
    trade = analysis.get("trade_name", "")
    inn = analysis.get("inn", "")
    
    _add_heading(doc, f"UAE 바이어 후보 리스트 — {trade}", level=0)
    _add_body(doc, f"UAE  |  {datetime.now().strftime('%Y-%m-%d')}")
    _add_body(doc, "※ 아래 바이어 후보는 CPHI 등록 및 Perplexity 웹 분석을 통해 도출되었으며, 개별 기업의 UAE 진출 현황 및 제품 연관성은 추가 실사가 필요합니다.", size=9)
    
    if not buyers:
        _add_body(doc, "바이어 분석을 먼저 실행해주세요.")
        doc.save(str(output_path))
        return output_path
    
    # 1. 전체 리스트
    _add_heading(doc, f"1. 바이어 후보 리스트 (전체 {len(buyers)}개사)", level=1)
    for i, b in enumerate(buyers, 1):
        name = b.get("company_name", "")
        country = b.get("country", "")
        category = b.get("category", "")
        email = b.get("email", "—")
        _add_body(doc, f"{i}.  {name}  |  {country}  |  {category}  |  {email}", size=10)
    
    # 2. 상위 3개사 상세
    top3 = buyers[:3]
    _add_heading(doc, f"2. 우선 접촉 바이어 상세 정보 (상위 {len(top3)}개사)", level=1)
    _add_body(doc, f"※ 하기 {len(top3)}개사는 {trade}의 성분과 연관성, UAE/중동 네트워크, 진출 가능성을 종합 평가하여 선정하였습니다.", size=9)
    
    for i, b in enumerate(top3, 1):
        name = b.get("company_name", "")
        country = b.get("country", "")
        category = b.get("category", "")
        e = b.get("enriched", {})
        
        _add_heading(doc, f"{i}. {name}  |  {country} · {category}", level=1)
        
        # 기업 개요
        overview = e.get("company_overview_kr", "")
        if overview:
            _add_heading(doc, "▸ 기업 개요", level=2)
            _add_body(doc, overview)
        
        # 추천 이유
        reason = e.get("recommendation_reason", "")
        if reason:
            _add_heading(doc, "▸ 추천 이유", level=2)
            _add_body(doc, reason)
        
        # 매출/파이프라인/제조소/수입경험/공급체인
        revenue = e.get("revenue", "")
        if revenue:
            _add_field(doc, "① 매출 규모", revenue)
        
        prods = b.get("products_cphi", [])
        if prods:
            _add_field(doc, "② 파이프라인", ", ".join(prods[:10]))
        
        gmp = e.get("has_gmp")
        if gmp is not None:
            _add_field(doc, "③ 제조소 보유", "GMP 인증 제조 시설 보유" if gmp else "제조소 미확인")
        
        imp = e.get("import_history")
        if imp is not None:
            _add_field(doc, "④ 수입 경험", "UAE/중동 의약품 수입·유통 경험 보유" if imp else "수입 이력 미확인")
        
        chain = e.get("has_pharmacy_chain")
        if chain is not None:
            _add_field(doc, "⑤ 약국 체인 유무", "약국 체인 보유/연계" if chain else "약국 체인 미확인")
        
        # 기본 연락처
        _add_field(doc, "주소", b.get("address", "—"))
        _add_field(doc, "이메일", b.get("email", "—"))
        _add_field(doc, "홈페이지", b.get("website", "—"))
        
        employees = e.get("employees", "")
        if employees:
            _add_field(doc, "기업 규모", employees)
        
        src_urls = e.get("source_urls", [])
        if src_urls:
            _add_body(doc, f"※ 출처: Perplexity 분석", size=8)
    
    doc.save(str(output_path))
    return output_path


# ── 합본 생성 ────────────────────────────────────────────────────────────────

def generate_all_reports(
    analysis: dict,
    buyers: list = None,
    p2_data: dict = None,
    output_dir: Path = None,
) -> dict:
    """4개 DOCX 파일 생성 후 경로 반환."""
    if output_dir is None:
        output_dir = ROOT / "reports" / "sg_format"
    output_dir.mkdir(parents=True, exist_ok=True)
    
    trade = analysis.get("trade_name", "product")
    safe_name = re.sub(r'[^\w가-힣]', '_', trade)[:20]
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    paths = {}
    
    # 00 표지
    p0 = output_dir / f"UAE_00_표지_{safe_name}_{ts}.docx"
    generate_cover(p0, trade)
    paths["cover"] = p0
    log.info(f"표지 생성: {p0}")
    
    # 01 시장보고서
    p1 = output_dir / f"UAE_01_시장보고서_{safe_name}_{ts}.docx"
    generate_market_report(p1, analysis)
    paths["market"] = p1
    log.info(f"시장보고서 생성: {p1}")
    
    # 02 수출가격전략
    p2 = output_dir / f"UAE_02_수출가격전략_{safe_name}_{ts}.docx"
    generate_price_strategy(p2, analysis, p2_data)
    paths["price"] = p2
    log.info(f"수출가격전략 생성: {p2}")
    
    # 03 바이어리스트
    p3 = output_dir / f"UAE_03_바이어리스트_{safe_name}_{ts}.docx"
    generate_buyer_list(p3, analysis, buyers)
    paths["buyer"] = p3
    log.info(f"바이어리스트 생성: {p3}")
    
    return paths


def convert_to_pdf(docx_paths: dict) -> dict:
    """DOCX 파일들을 PDF로 변환."""
    from docx2pdf import convert
    pdf_paths = {}
    for key, docx_path in docx_paths.items():
        pdf_path = docx_path.with_suffix('.pdf')
        try:
            convert(str(docx_path), str(pdf_path))
            pdf_paths[key] = pdf_path
            log.info(f"PDF 변환 완료: {pdf_path}")
        except Exception as e:
            log.error(f"PDF 변환 실패 ({key}): {e}")
    return pdf_paths


# ── CLI 테스트 ────────────────────────────────────────────────────────────────

async def _test_with_real_data():
    """실제 Supabase 데이터로 테스트."""
    from analysis.uae_export_analyzer import analyze_custom_product
    
    print("Seretide 분석 중...")
    result = await analyze_custom_product("Seretide", "Fluticasone/Salmeterol", "Inhaler")
    
    print(f"분석 완료: verdict={result.get('verdict')}")
    print(f"price_comparison: {len(result.get('price_comparison', {}).get('competitors', []))}건")
    
    paths = generate_all_reports(
        analysis=result,
        buyers=None,
        p2_data=None,
    )
    
    print("\n생성된 파일:")
    for key, path in paths.items():
        print(f"  {key}: {path}")
    
    # PDF 변환 시도
    try:
        pdf_paths = convert_to_pdf(paths)
        print("\nPDF 변환 완료:")
        for key, path in pdf_paths.items():
            print(f"  {key}: {path}")
    except Exception as e:
        print(f"\nPDF 변환 실패 (Word 미설치?): {e}")


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    asyncio.run(_test_with_real_data())
